"""文档解析与切分服务

功能：
1. 解析多种格式文档（PDF、TXT、Markdown、DOCX）
2. 文档切分（支持递归切分、父子文档切分）
3. 元数据提取
"""
from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import yaml


class DocumentProcessor:
    """文档处理器 - 解析 + 切分"""

    def __init__(self, config_path: str = None):
        """初始化文档处理器

        Args:
            config_path: 配置文件路径，默认使用项目根目录下的 config/config.yaml
        """
        if config_path is None:
            config_path = Path(__file__).parent.parent.parent / "config" / "config.yaml"

        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = yaml.safe_load(f)

        self.chunking_config = self.config.get('chunking', {})

    def parse_file(self, file_path: str) -> str:
        """解析单个文件，提取纯文本内容

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.pdf':
            return self._parse_pdf(file_path)
        elif suffix == '.txt':
            return self._parse_txt(file_path)
        elif suffix == '.md':
            return self._parse_txt(file_path)  # Markdown 当纯文本处理
        elif suffix == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{suffix}")

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件"""
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本 / Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _parse_docx(self, file_path: str) -> str:
        """解析 Word 文档"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n\n".join(paragraphs)

    def parse_bytes(self, file_bytes: bytes, filename: str) -> str:
        """从字节流解析文件（用于 API 上传场景）

        Args:
            file_bytes: 文件字节内容
            filename: 文件名（用于判断格式）

        Returns:
            提取的文本内容
        """
        import tempfile

        # 写入临时文件再解析
        suffix = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            return self.parse_file(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def chunk_documents(self, text: str, source: str = "unknown",
                        strategy: str = None,
                        chunk_size: int = None,
                        chunk_overlap: int = None) -> List[Document]:
        """将文本切分为文档块

        切分策略：
        - recursive: 递归字符切分，块大小可调，直接检索该块
        - parent_child: 父子切分，只索引小子块、命中后返回大父块（small-to-big）

        Args:
            text: 要切分的文本
            source: 来源文件名（写入元数据）
            strategy: 覆盖配置中的切分策略（None 则用 config）
            chunk_size: 覆盖块大小（parent_child 时代表父块大小）
            chunk_overlap: 覆盖重叠大小（parent_child 时代表父块重叠）

        Returns:
            切分后的 Document 列表
        """
        strategy = strategy or self.chunking_config.get('strategy', 'recursive')

        # 前端传入的 chunk_size/overlap 覆盖父块(或递归块)默认值
        parent_config = self.chunking_config.get('parent', {})
        p_size = chunk_size if chunk_size is not None else parent_config.get('chunk_size', 800)
        p_overlap = chunk_overlap if chunk_overlap is not None else parent_config.get('chunk_overlap', 200)

        if strategy == 'parent_child':
            return self._parent_child_chunk(text, source, p_size, p_overlap)
        else:
            return self._recursive_chunk(text, source, p_size, p_overlap)

    def _recursive_chunk(self, text: str, source: str,
                         chunk_size: int = None,
                         chunk_overlap: int = None) -> List[Document]:
        """递归字符切分"""
        parent_config = self.chunking_config.get('parent', {})
        if chunk_size is None:
            chunk_size = parent_config.get('chunk_size', 800)
        if chunk_overlap is None:
            chunk_overlap = parent_config.get('chunk_overlap', 200)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", " ", ""]
        )

        chunks = splitter.split_text(text)
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_id": i,
                    "total_chunks": len(chunks),
                    "chunk_type": "parent"
                }
            )
            documents.append(doc)

        return documents

    def _parent_child_chunk(self, text: str, source: str,
                            parent_size: int, parent_overlap: int) -> List[Document]:
        """父子文档切分（small-to-big）

        只把【子块】作为可检索单元入库：子块小、语义聚焦，向量/BM25 命中更准。
        每个子块的 metadata 携带其父块全文 parent_content，检索命中子块后
        由 rag_chain 展开回父块交给 LLM，兼顾"检索精度"和"上下文完整"。

        Args:
            text: 要切分的文本
            source: 来源文件名
            parent_size: 父块大小（检索返回的上下文单元）
            parent_overlap: 父块重叠

        Returns:
            子块 Document 列表（每个带 parent_content 元数据）
        """
        child_config = self.chunking_config.get('child', {})
        child_size = child_config.get('chunk_size', 200)
        child_overlap = child_config.get('chunk_overlap', 50)

        separators = ["\n\n", "\n", "。", "！", "？", ".", " ", ""]

        # 第一步：切大父块
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=parent_size,
            chunk_overlap=parent_overlap,
            separators=separators,
        )
        parent_chunks = parent_splitter.split_text(text)

        # 第二步：每个父块切小子块，子块携带父块全文
        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=child_size,
            chunk_overlap=child_overlap,
            separators=separators,
        )

        documents = []
        for parent_id, parent_text in enumerate(parent_chunks):
            child_chunks = child_splitter.split_text(parent_text)
            for child_id, child_text in enumerate(child_chunks):
                documents.append(Document(
                    page_content=child_text,          # 用于检索的是子块
                    metadata={
                        "source": source,
                        "parent_id": parent_id,
                        "child_id": child_id,
                        "chunk_type": "child",
                        "doc_id": f"{source}_parent_{parent_id}_child_{child_id}",
                        # 检索命中后展开回这个父块
                        "parent_content": parent_text,
                        "parent_doc_id": f"{source}_parent_{parent_id}",
                    }
                ))

        return documents

    def process_file(self, file_path: str, strategy: str = None,
                     chunk_size: int = None, chunk_overlap: int = None) -> List[Document]:
        """完整处理流程：解析 + 切分

        Args:
            file_path: 文件路径
            strategy / chunk_size / chunk_overlap: 覆盖配置的切分参数

        Returns:
            处理后的 Document 列表
        """
        # 解析文件
        text = self.parse_file(file_path)
        if not text.strip():
            return []

        # 切分文档
        source = Path(file_path).name
        documents = self.chunk_documents(text, source, strategy, chunk_size, chunk_overlap)

        return documents

    def process_bytes(self, file_bytes: bytes, filename: str, strategy: str = None,
                      chunk_size: int = None, chunk_overlap: int = None) -> List[Document]:
        """完整处理流程（字节流版本）：解析 + 切分

        Args:
            file_bytes: 文件字节内容
            filename: 文件名

        Returns:
            处理后的 Document 列表
        """
        text = self.parse_bytes(file_bytes, filename)
        if not text.strip():
            return []

        documents = self.chunk_documents(text, filename, strategy, chunk_size, chunk_overlap)
        return documents
